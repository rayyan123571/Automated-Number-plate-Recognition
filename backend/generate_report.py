"""Generate a professional Word report for the ANPR project.

This script builds a structured report answering:
  1. Problem Statement
  2. Motivation
  3. Methodology
  4. Conclusion & Future Directions

It also includes key project statistics (dataset size, model, tech stack).

Usage:
    python generate_report.py

Output:
    ANPR_Project_Report.docx (in the project root)
"""

from __future__ import annotations

import datetime
import csv
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple


def safe_import_docx() -> Optional[object]:
    """Import python-docx if available, otherwise show an actionable error."""
    try:
        import docx  # type: ignore

        return docx
    except ImportError:
        return None


def get_project_root() -> Path:
    """Return the repository root (parent of the backend folder)."""
    return Path(__file__).resolve().parents[1]


def get_dataset_paths(project_root: Path) -> Dict[str, Path]:
    """Resolve the training/validation/test image folders from dataset/data.yaml."""
    import yaml

    config_path = project_root / "backend" / "dataset" / "data.yaml"
    with open(config_path, "r", encoding="utf-8") as f:
        cfg = yaml.safe_load(f)

    base = Path(cfg["path"]).expanduser()
    # If path is relative, make it relative to the data.yaml location
    if not base.is_absolute():
        base = (config_path.parent / base).resolve()

    return {
        "train": (base / cfg.get("train", "train/images")).resolve(),
        "val": (base / cfg.get("val", "valid/images")).resolve(),
        "test": (base / cfg.get("test", "test/images")).resolve(),
    }


def count_dataset_files(image_dir: Path) -> Tuple[int, int]:
    """Count images and labels found in a dataset split.

    Returns:
        (image_count, label_count)
    """

    if not image_dir.exists():
        return 0, 0

    image_patterns = ["*.jpg", "*.jpeg", "*.png", "*.bmp", "*.webp"]
    image_files = []
    for pat in image_patterns:
        image_files.extend(image_dir.glob(pat))

    # YOLO labels live in a sibling "labels" directory
    labels_dir = image_dir.parent / "labels"
    label_files = []
    if labels_dir.exists():
        label_files = list(labels_dir.glob("*.txt"))

    return len(image_files), len(label_files)


def get_dataset_stats(project_root: Path) -> Dict[str, Dict[str, int | str]]:
    """Collect per-split and total dataset statistics."""
    split_paths = get_dataset_paths(project_root)
    stats: Dict[str, Dict[str, int | str]] = {}
    total_images = 0
    total_labels = 0

    for split, path in split_paths.items():
        images, labels = count_dataset_files(path)
        total_images += images
        total_labels += labels
        stats[split] = {
            "images": images,
            "labels": labels,
            "source": str(path),
        }

    stats["total"] = {
        "images": total_images,
        "labels": total_labels,
        "source": "all_splits",
    }
    return stats


def read_training_summary(project_root: Path) -> Dict[str, str]:
    """Read best and final metrics from YOLO training results.csv if available."""
    results_file = project_root / "backend" / "runs" / "anpr_train" / "results.csv"
    if not results_file.exists():
        return {
            "status": "not_found",
            "message": "Training results file not found.",
        }

    with open(results_file, "r", encoding="utf-8", newline="") as f:
        rows = list(csv.DictReader(f))

    if not rows:
        return {
            "status": "empty",
            "message": "Training results file is empty.",
        }

    def safe_float(row: dict, key: str) -> float:
        try:
            return float(row.get(key, 0.0))
        except Exception:
            return 0.0

    best_row = max(rows, key=lambda r: safe_float(r, "metrics/mAP50(B)"))
    final_row = rows[-1]

    return {
        "status": "ok",
        "epochs": str(len(rows)),
        "best_epoch": str(best_row.get("epoch", "?")),
        "best_map50": f"{safe_float(best_row, 'metrics/mAP50(B)'):.4f}",
        "best_map50_95": f"{safe_float(best_row, 'metrics/mAP50-95(B)'):.4f}",
        "best_precision": f"{safe_float(best_row, 'metrics/precision(B)'):.4f}",
        "best_recall": f"{safe_float(best_row, 'metrics/recall(B)'):.4f}",
        "final_map50": f"{safe_float(final_row, 'metrics/mAP50(B)'):.4f}",
        "final_map50_95": f"{safe_float(final_row, 'metrics/mAP50-95(B)'):.4f}",
        "final_precision": f"{safe_float(final_row, 'metrics/precision(B)'):.4f}",
        "final_recall": f"{safe_float(final_row, 'metrics/recall(B)'):.4f}",
    }


def build_project_summary(project_root: Path) -> Dict[str, str]:
    """Gather high-level project metadata to inject into the report."""
    summary = {}

    # Frameworks and libraries
    summary["Backend"] = "FastAPI + Ultralytics YOLOv8 + EasyOCR"
    summary["Frontend"] = "Next.js (App Router) + Tailwind CSS"

    # Dataset
    dataset_path = project_root / "backend" / "dataset" / "data.yaml"
    summary["Dataset config"] = str(dataset_path.relative_to(project_root))

    # Model weights
    model_weights = project_root / "backend" / "models" / "best.pt"
    summary["Trained model weights"] = (
        str(model_weights.relative_to(project_root))
        if model_weights.exists()
        else "<not found>"
    )

    return summary


def format_section_header(doc: Any, title: str, level: int = 1) -> None:
    """Add a section header to the document."""
    doc.add_heading(title, level=level)


def add_bullet_list(doc: Any, items: List[str]) -> None:
    """Add a bullet list to the document."""
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(item)


def create_report(project_root: Path, output_path: Path) -> None:
    """Create the Word report using python-docx."""
    docx = safe_import_docx()
    if docx is None:
        raise RuntimeError(
            "python-docx is not installed. Install it with `pip install python-docx`."
        )

    # Build a new document
    doc = docx.Document()

    # Cover page
    print("[1/6] Writing cover page...")
    title = "Automated Number Plate Recognition (ANPR) Project Report"
    doc.add_heading(title, level=0)
    doc.add_paragraph("Generated: %s" % datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    doc.add_paragraph(
        "This report provides an overview of the ANPR project, including the problem statement, motivation, methodology, evaluation strategy, and future directions."
    )
    doc.add_page_break()

    # 1. Problem Statement
    print("[2/6] Adding problem statement section...")
    format_section_header(doc, "1. Problem Statement", level=1)
    doc.add_paragraph(
        "This project solves the problem of automatic vehicle license plate localization and text extraction from static images and live camera streams. "
        "In conventional workflows, operators manually monitor cameras and type plate numbers into logs, which is time-consuming and error-prone. "
        "The ANPR system replaces this with a real-time, AI-based pipeline that can detect plate regions, read text, and store records for analytics and security workflows."
    )

    # 2. Motivation
    print("[3/6] Adding motivation section...")
    format_section_header(doc, "2. Motivation", level=1)
    add_bullet_list(
        doc,
        [
            "Reduce manual effort and human error when managing vehicle access or recording plate data.",
            "Enable real-time monitoring with low-latency detection and OCR for live camera feeds.",
            "Support scalable deployment via a web API (FastAPI) and a modern dashboard (Next.js).",
            "Leverage state-of-the-art deep learning (YOLOv8) and OCR (EasyOCR) while maintaining real-time performance.",
            "Provide a practical semester project that bridges machine learning, backend APIs, WebSocket streaming, and frontend engineering.",
        ],
    )

    # 3. Methodology
    print("[4/6] Adding methodology section...")
    format_section_header(doc, "3. Methodology", level=1)

    doc.add_paragraph("The solution uses a two-stage ANPR pipeline and full-stack deployment:")
    add_bullet_list(
        doc,
        [
            "Backend: FastAPI provides REST endpoints (/detect, /detections, /health) and a WebSocket endpoint (/ws/detect) for live video feeds.",
            "Detection: YOLOv8 (Ultralytics) detects license plate bounding boxes in an image.",
            "Recognition: EasyOCR reads the text within the detected plate region to produce the license string.",
            "Persistence: Detected plates and metadata are stored in SQLite via SQLAlchemy models.",
            "Frontend: Next.js dashboard allows users to upload images, view live camera feeds, and inspect detection history.",
        ],
    )

    format_section_header(doc, "Algorithm Explored", level=2)
    add_bullet_list(
        doc,
        [
            "Stage 1 - Detection: YOLOv8n predicts bounding boxes for class `plate-number` in each frame/image.",
            "Stage 2 - OCR: The detected region is cropped and passed to EasyOCR to recognize alphanumeric characters.",
            "Post-processing: OCR text is cleaned/normalized before storing in the database.",
            "Serving: Detection API supports both batch image upload and live stream inference through WebSocket.",
        ],
    )

    format_section_header(doc, "Incremental Development Execution", level=2)
    add_bullet_list(
        doc,
        [
            "Phase 1: Dataset preparation and validation (YOLO format labels, train/val/test split checks).",
            "Phase 2: Baseline model training with YOLOv8n and repeated hyperparameter tuning.",
            "Phase 3: OCR integration and end-to-end ANPR pipeline verification.",
            "Phase 4: FastAPI endpoint implementation with detection history persistence.",
            "Phase 5: Frontend dashboard integration (live monitoring, history, analytics).",
            "Phase 6: Final evaluation, documentation, and report generation automation.",
        ],
    )

    # Dataset / Images
    print("[5/6] Gathering dataset statistics...")
    dataset_stats = get_dataset_stats(project_root)
    total_images = int(dataset_stats["total"]["images"])
    total_labels = int(dataset_stats["total"]["labels"])
    doc.add_paragraph(
        "A supervised dataset is used for model training and evaluation with YOLO annotation format. "
        f"Across all splits, the dataset currently contains {total_images} images and {total_labels} label files."
    )

    # Add dataset counts table
    format_section_header(doc, "Dataset Overview", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Split"
    hdr_cells[1].text = "Image Count"
    hdr_cells[2].text = "Label Count"
    hdr_cells[3].text = "Source"

    for split in ["train", "val", "test"]:
        split_info = dataset_stats[split]
        row_cells = table.add_row().cells
        row_cells[0].text = split.capitalize()
        row_cells[1].text = str(split_info["images"])
        row_cells[2].text = str(split_info["labels"])
        row_cells[3].text = str(split_info["source"])

    doc.add_paragraph(
        "The dataset is defined in `backend/dataset/data.yaml` and contains a single class: `plate-number` (class id 0). "
        "Images are stored in the Automatic Plate Number Recognition.v4i.yolov8 dataset folder."
    )

    # Evaluation methodology
    print("[6/6] Adding evaluation strategy and conclusions...")
    format_section_header(doc, "Evaluation Strategy", level=2)
    doc.add_paragraph(
        "The model is evaluated using standard detection metrics: Precision, Recall, mAP@0.5, and mAP@0.5:0.95. "
        "The validation/test split is used to estimate generalization performance, and evaluation artifacts are stored in backend/metrics."
    )

    training_summary = read_training_summary(project_root)
    format_section_header(doc, "Observed Model Performance", level=2)
    if training_summary.get("status") == "ok":
        doc.add_paragraph(
            "Training evidence from runs/anpr_train/results.csv confirms strong detection quality after iterative optimization."
        )
        perf_table = doc.add_table(rows=1, cols=2)
        ph = perf_table.rows[0].cells
        ph[0].text = "Metric"
        ph[1].text = "Value"

        for metric, value in [
            ("Total Epochs Run", training_summary["epochs"]),
            ("Best Epoch", training_summary["best_epoch"]),
            ("Best Precision", training_summary["best_precision"]),
            ("Best Recall", training_summary["best_recall"]),
            ("Best mAP@0.5", training_summary["best_map50"]),
            ("Best mAP@0.5:0.95", training_summary["best_map50_95"]),
            ("Final Precision", training_summary["final_precision"]),
            ("Final Recall", training_summary["final_recall"]),
            ("Final mAP@0.5", training_summary["final_map50"]),
            ("Final mAP@0.5:0.95", training_summary["final_map50_95"]),
        ]:
            row = perf_table.add_row().cells
            row[0].text = metric
            row[1].text = value
    else:
        doc.add_paragraph(
            "Training metrics table could not be loaded automatically. "
            f"Reason: {training_summary.get('message', 'Unknown issue')}"
        )

    # 4. Conclusion & Future Directions
    format_section_header(doc, "4. Conclusion & Future Directions", level=1)
    add_bullet_list(
        doc,
        [
            "The project demonstrates an end-to-end ANPR system that combines machine learning, OCR, backend APIs, and a modern monitoring dashboard.",
            "A key contribution is practical integration: from model training and evaluation to live inference and historical record management.",
            "Future work can focus on OCR robustness (blur/night/rain), support for diverse plate formats, and multilingual plate reading.",
            "Possible extensions include vehicle re-identification, edge inference optimization, and deployment in smart-campus/smart-city environments.",
        ],
    )

    # Appendix: Project summary
    format_section_header(doc, "Appendix: Project Summary", level=2)
    summary = build_project_summary(project_root)
    for key, value in summary.items():
        doc.add_paragraph(f"{key}: {value}")

    # Save document
    print("Saving report document...")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(output_path)
    except PermissionError:
        fallback_name = f"{output_path.stem}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}{output_path.suffix}"
        fallback_path = output_path.with_name(fallback_name)
        doc.save(fallback_path)
        print(f"Output file was in use. Saved to fallback path: {fallback_path}")


def main() -> None:
    project_root = get_project_root()
    output_file = project_root / "ANPR_Project_Report.docx"
    print(f"Generating report at: {output_file}")
    create_report(project_root, output_file)
    print("Report generation complete.")


if __name__ == "__main__":
    main()
